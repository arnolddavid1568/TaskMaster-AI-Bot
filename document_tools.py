"""
Conversions between TXT, DOCX, and PDF.
Uses python-docx for DOCX and reportlab for generating PDFs from text,
plus pypdf for reading PDF text back out.
"""
from pathlib import Path

import docx
from pypdf import PdfReader
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas


def txt_to_pdf(txt_path: str, output_path: str) -> str:
    text = Path(txt_path).read_text(encoding="utf-8", errors="replace")
    c = canvas.Canvas(output_path, pagesize=A4)
    width, height = A4
    margin = 20 * mm
    y = height - margin
    line_height = 14

    c.setFont("Helvetica", 11)
    for raw_line in text.splitlines() or [""]:
        # naive wrap at ~95 chars for Helvetica 11 on A4
        for line in _wrap(raw_line, 95):
            if y < margin:
                c.showPage()
                c.setFont("Helvetica", 11)
                y = height - margin
            c.drawString(margin, y, line)
            y -= line_height
    c.save()
    return output_path


def _wrap(line: str, width: int):
    if not line:
        return [""]
    out = []
    while len(line) > width:
        out.append(line[:width])
        line = line[width:]
    out.append(line)
    return out


def docx_to_txt(docx_path: str, output_path: str) -> str:
    d = docx.Document(docx_path)
    text = "\n".join(p.text for p in d.paragraphs)
    Path(output_path).write_text(text, encoding="utf-8")
    return output_path


def txt_to_docx(txt_path: str, output_path: str) -> str:
    text = Path(txt_path).read_text(encoding="utf-8", errors="replace")
    d = docx.Document()
    for line in text.splitlines() or [""]:
        d.add_paragraph(line)
    d.save(output_path)
    return output_path


def pdf_to_txt(pdf_path: str, output_path: str) -> str:
    reader = PdfReader(pdf_path)
    text = "\n\n".join((page.extract_text() or "") for page in reader.pages)
    Path(output_path).write_text(text, encoding="utf-8")
    return output_path


def docx_to_pdf(docx_path: str, output_path: str) -> str:
    """Best-effort DOCX -> PDF: extracts paragraph text and lays it out
    with reportlab (keeps text content; does not preserve rich styling)."""
    d = docx.Document(docx_path)
    c = canvas.Canvas(output_path, pagesize=A4)
    width, height = A4
    margin = 20 * mm
    y = height - margin
    line_height = 14

    c.setFont("Helvetica", 11)
    for para in d.paragraphs:
        for line in _wrap(para.text, 95) or [""]:
            if y < margin:
                c.showPage()
                c.setFont("Helvetica", 11)
                y = height - margin
            c.drawString(margin, y, line)
            y -= line_height
    c.save()
    return output_path
